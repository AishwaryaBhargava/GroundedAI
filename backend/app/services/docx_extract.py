from typing import List
import re
import io
from docx import Document


def _normalize_text(s: str) -> str:
    s = s.replace("\u00ad", "")  # soft hyphen
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def extract_pages(docx_bytes: bytes) -> List[dict]:
    """
    Extract DOCX content into pseudo-pages.

    Uses word-count based paging to stay compatible
    with PDF-style page abstraction.
    """

    # 🔹 FIX: wrap bytes in file-like object
    file_like = io.BytesIO(docx_bytes)
    doc = Document(file_like)

    blocks: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    pages: List[dict] = []
    buffer: List[str] = []
    word_count = 0
    page_num = 1

    for block in blocks:
        words = block.split()
        buffer.append(block)
        word_count += len(words)

        if word_count >= 900:
            pages.append(
                {
                    "page": page_num,
                    "text": _normalize_text("\n".join(buffer)),
                }
            )
            page_num += 1
            buffer = []
            word_count = 0

    # Flush remainder
    if buffer:
        pages.append(
            {
                "page": page_num,
                "text": _normalize_text("\n".join(buffer)),
            }
        )

    return pages
